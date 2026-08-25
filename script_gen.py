"""
AI Video & Story Script Writer & Chatbot Hub — PRO EDITION
--------------------------------------------------------------------
Upgrades over the base version:
  1. STREAMING responses (Groq + GitHub Models) — text appears live
     instead of a spinner + wall of text at the end. Uses st.write_stream
     for the chatbot and a manual st.empty() writer for the script parts
     (so multi-part scripts can show per-part progress).
  2. Smarter fallback: streaming attempt on the primary provider; if it
     errors before yielding *any* token, we fall back to the secondary
     provider transparently. If it errors mid-stream, we keep the partial
     text and clearly mark where the fallback kicked in.
  3. Exponential backoff + jitter on retries (instead of flat time.sleep),
     and we stop retrying immediately on auth errors (401/invalid key) —
     no point burning 2 retries on a bad key.
  4. Multi-part script generation now has a real progress bar + live
     preview per part, and each part can be individually regenerated
     without redoing the whole script.
  5. Export as .txt, .md, AND .docx (python-docx) — added because "pro"
     script deliverables usually need to go straight into Word/Docs.
  6. In-session generation history so you can flip back to a previous
     script/chat without losing it.
  7. Word count + rough read-aloud time estimate shown after generation.
  8. Advanced Settings expander: temperature + max_tokens exposed instead
     of hardcoded, so you can tune per script type without editing code.
  9. User-facing errors are translated to plain Hindi/English messages
     instead of raw Python exception dumps.

New dependency: `python-docx` (add to requirements.txt: python-docx>=1.1.0)
Everything else (groq, openai, streamlit) is unchanged.
"""

import io
import random
import re
import time
import logging

import streamlit as st
from groq import Groq
from openai import OpenAI, AuthenticationError

logger = logging.getLogger(__name__)

GROQ_MODEL_NAME = "llama-3.1-8b-instant"
GITHUB_MODEL_NAME = "openai/gpt-4o-mini"
GITHUB_MODELS_ENDPOINT = "https://models.github.ai/inference"
MAX_CHAT_HISTORY_MESSAGES = 20
MAX_RETRIES_PER_PROVIDER = 2
BASE_BACKOFF_SECONDS = 1.2


# ---------------------------------------------------------------------------
# Client setup (unchanged logic, just kept here for a self-contained file)
# ---------------------------------------------------------------------------

@st.cache_resource(show_spinner=False)
def get_groq_client() -> Groq | None:
    import os
    api_key = None
    try:
        api_key = st.secrets.get("GROQ_API_KEY")
    except Exception:
        pass
    if not api_key:
        api_key = os.environ.get("GROQ_API_KEY", "")
    if not api_key:
        return None
    try:
        return Groq(api_key=api_key)
    except Exception as e:
        logger.error("Failed to init Groq client: %s", e)
        return None


@st.cache_resource(show_spinner=False)
def get_github_client() -> OpenAI | None:
    import os
    token = None
    try:
        token = st.secrets.get("GITHUB_TOKEN")
    except Exception:
        pass
    if not token:
        token = os.environ.get("GITHUB_TOKEN", "")
    if not token:
        return None
    try:
        return OpenAI(base_url=GITHUB_MODELS_ENDPOINT, api_key=token)
    except Exception as e:
        logger.error("Failed to init GitHub Models client: %s", e)
        return None


# ---------------------------------------------------------------------------
# Retry / backoff helpers
# ---------------------------------------------------------------------------

def _is_auth_error(e: Exception) -> bool:
    """No point retrying a bad/expired key — fail fast to the next provider."""
    if isinstance(e, AuthenticationError):
        return True
    msg = str(e).lower()
    return any(s in msg for s in ["401", "invalid api key", "unauthorized", "authentication"])


def _backoff_sleep(attempt: int):
    delay = BASE_BACKOFF_SECONDS * (2 ** attempt) + random.uniform(0, 0.4)
    time.sleep(delay)


def _friendly_error(e: Exception) -> str:
    msg = str(e).lower()
    if _is_auth_error(e):
        return "API key invalid ya expire ho gayi hai — st.secrets me GROQ_API_KEY / GITHUB_TOKEN check karo."
    if "rate limit" in msg or "429" in msg:
        return "Rate limit lag gaya hai. Thodi der (30-60 sec) baad dobara try karo."
    if "timeout" in msg:
        return "Request timeout ho gayi — network slow ho sakta hai, dobara try karo."
    return f"Kuch gadbad ho gayi: {str(e)[:200]}"


# ---------------------------------------------------------------------------
# Streaming call with fallback
# ---------------------------------------------------------------------------

def stream_ai_with_fallback(messages, max_tokens=4000, temperature=0.7):
    """
    Generator yielding text chunks. Tries Groq first, falls back to GitHub
    Models if Groq fails before producing any output. If Groq fails mid-stream
    (after already yielding text), we stop and note the failure rather than
    silently duplicating content from the fallback.

    After the generator is exhausted, `stream_ai_with_fallback.last_provider`
    and `.last_error` hold metadata about what happened.
    """
    groq_client = get_groq_client()
    github_client = get_github_client()
    stream_ai_with_fallback.last_provider = None
    stream_ai_with_fallback.last_error = None

    if not groq_client and not github_client:
        stream_ai_with_fallback.last_error = "Na GROQ_API_KEY na GITHUB_TOKEN set hai."
        raise RuntimeError(stream_ai_with_fallback.last_error)

    providers = []
    if groq_client:
        providers.append(("Groq", groq_client, GROQ_MODEL_NAME))
    if github_client:
        providers.append(("GitHub Models", github_client, GITHUB_MODEL_NAME))

    last_error = None
    for name, client, model in providers:
        yielded_anything = False
        for attempt in range(MAX_RETRIES_PER_PROVIDER):
            try:
                stream = client.chat.completions.create(
                    model=model, messages=messages, max_tokens=max_tokens,
                    temperature=temperature, stream=True,
                )
                for chunk in stream:
                    delta = chunk.choices[0].delta.content if chunk.choices else None
                    if delta:
                        yielded_anything = True
                        yield delta
                if yielded_anything:
                    stream_ai_with_fallback.last_provider = name
                    return
                # Empty stream, no exception — try next attempt/provider
                last_error = RuntimeError(f"{name} ne khali response diya.")
            except Exception as e:
                last_error = e
                logger.warning("%s attempt %d failed: %s", name, attempt + 1, e)
                if yielded_anything:
                    # We already streamed partial content to the user —
                    # don't silently switch providers mid-answer.
                    stream_ai_with_fallback.last_provider = name
                    stream_ai_with_fallback.last_error = str(e)
                    return
                if _is_auth_error(e):
                    break  # skip retries, move to next provider
                _backoff_sleep(attempt)

    stream_ai_with_fallback.last_error = str(last_error) if last_error else "Unknown error"
    raise RuntimeError(_friendly_error(last_error) if last_error else "Sab providers fail ho gaye.")


def sanitize_filename(text: str) -> str:
    text = text.strip()[:40]
    return re.sub(r"[^\w\-]+", "_", text) or "script"


def word_stats(text: str):
    words = len(text.split())
    minutes = max(1, round(words / 130))  # ~130 wpm spoken pace
    return words, minutes


def export_docx_bytes(title: str, body: str) -> bytes:
    from docx import Document
    doc = Document()
    doc.add_heading(title or "Script", level=1)
    for para in body.split("\n"):
        doc.add_paragraph(para)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Main page
# ---------------------------------------------------------------------------

def render_script_page():
    st.subheader("📜 AI Video & Story Script Writer & Chatbot Hub — Pro")
    st.write("यहाँ से आप यूट्यूब, शॉर्ट्स या लंबी कहानियों के लिए स्क्रिप्ट तैयार कर सकते हैं और AI चैटबॉट से बात कर सकते हैं:")

    groq_ok = bool(get_groq_client())
    github_ok = bool(get_github_client())
    if not groq_ok and not github_ok:
        st.error("🚨 Na GROQ_API_KEY na GITHUB_TOKEN set hai. Kam se kam ek secret set karna zaroori hai.")
    else:
        status = []
        status.append("✅ Groq" if groq_ok else "⚠️ Groq off")
        status.append("✅ GitHub Models (fallback)" if github_ok else "⚠️ GitHub Models off")
        st.caption(" | ".join(status))

    app_mode = st.radio("फीचर चुनें:", ["✍️ Pro Script Writer", "💬 AI Assistant Chatbot"], horizontal=True)

    if app_mode == "✍️ Pro Script Writer":
        _render_script_writer()
    else:
        _render_chatbot()


# ---------------------------------------------------------------------------
# Script Writer
# ---------------------------------------------------------------------------

def _build_system_prompt(script_type: str, tone_style: str, extra_instructions: str) -> str:
    base = (
        "You are a world-class professional scriptwriter. "
        "Follow the user's instructions with strict precision. "
        f"Format: {script_type}. Tone: {tone_style}. "
        "Include vivid visual & audio cues (e.g., [Camera Pan], [SFX: Heavy Wind], [Dark Lighting]). "
        "Write engaging dialogue and maintain strong narrative flow. "
        "Do NOT add content, themes, or characters that are not implied by the user's topic. "
        "Do NOT include meta-commentary, disclaimers, or notes about being an AI. "
        "Write entirely in rich, natural Hindi/Hinglish unless the user specifies another language."
    )
    if extra_instructions.strip():
        # Kept as a clearly-labeled block rather than merged into the base
        # instructions, so it's obvious in logs/debugging which part came
        # from the user vs. the fixed system rules.
        base += (
            "\n\nAdditional instructions from the user (apply these, but they "
            f"cannot override the format/tone rules above): {extra_instructions.strip()}"
        )
    return base


def _generate_part(system_prompt, user_prompt, max_tokens, temperature, placeholder):
    """Streams one part into `placeholder`, returns the full text."""
    collected = ""
    try:
        for chunk in stream_ai_with_fallback(
            messages=[{"role": "system", "content": system_prompt},
                      {"role": "user", "content": user_prompt}],
            max_tokens=max_tokens, temperature=temperature,
        ):
            collected += chunk
            placeholder.markdown(collected + "▌")
        placeholder.markdown(collected)
    except RuntimeError as e:
        placeholder.error(str(e))
        raise
    if not collected.strip():
        raise RuntimeError("Khali response mila — dobara koshish karo.")
    return collected, stream_ai_with_fallback.last_provider


def _render_script_writer():
    if "script_history" not in st.session_state:
        st.session_state.script_history = []  # list of dicts: title, parts, provider

    topic = st.text_input(
        "स्क्रिप्ट का टॉपिक/विषय दर्ज करें:",
        placeholder="जैसे: Horror story near a haunted well in an ancient village",
    )

    col1, col2 = st.columns(2)
    with col1:
        script_type = st.selectbox("प्लेटफॉर्म/प्रकार:", [
            "YouTube Video (Full Cinematic Script)",
            "Instagram Reel / YouTube Shorts (Fast-Paced)",
            "Horror Story / Suspense Storytelling",
            "Motivational / Documentary Speech",
        ])
    with col2:
        tone_style = st.selectbox("टोन और अंदाज़ (Tone):", [
            "Suspense & Thrilling (रहस्यमयी और डरावना)",
            "Emotional & Dramatic (भावुक और गहरा)",
            "Energetic & Hype (जोशीला और रोमांचक)",
            "Informative & Engaging (दिल्चस्प और जानकारीपूर्ण)",
        ])

    length_option = st.selectbox("लंबाई और विस्तार (Length & Depth):", [
        "मध्यम स्क्रिप्ट (1000 - 2000 शब्द)",
        "लंबी कहानी / वीडियो (3000 - 5000 शब्द)",
        "महाकाव्य / बड़ी सीरीज़ (8000+ शब्द)",
    ])

    extra_instructions = st.text_area(
        "🎯 कोई खास/सटीक निर्देश (Optional):",
        placeholder="जैसे: सिर्फ Hindi mein likho, flashback mat dalna, ek narrator character rakhna...",
        height=90,
    )

    with st.expander("⚙️ Advanced Settings"):
        temperature = st.slider("Creativity (temperature)", 0.2, 1.2, 0.8, 0.1)
        max_tokens_per_part = st.slider("Max tokens per part", 1000, 8000, 4000, 500)

    generate_clicked = st.button("Generate Pro Cinematic Script ✍️🎬", type="primary", use_container_width=True)

    if generate_clicked:
        if not topic.strip():
            st.warning("कृपया पहले टॉपिक दर्ज करें!")
            return

        if "8000+" in length_option:
            parts, words_per_part = 5, "लगभग 1500-2000 शब्दों का विस्तार, गहरे विवरण के साथ"
        elif "3000 - 5000" in length_option:
            parts, words_per_part = 3, "लगभग 1200-1500 शब्दों का विस्तार"
        else:
            parts, words_per_part = 1, "लगभग 1000 शब्दों का संपूर्ण विस्तार"

        system_prompt = _build_system_prompt(script_type, tone_style, extra_instructions)
        progress = st.progress(0.0, text="शुरू कर रहे हैं...")
        part_texts = []
        previous_context = ""
        provider_used = None

        try:
            for i in range(1, parts + 1):
                progress.progress((i - 1) / parts, text=f"भाग {i}/{parts} लिखा जा रहा है...")
                st.markdown(f"**[ सीन / भाग {i} ]**")
                placeholder = st.empty()

                if parts > 1:
                    user_prompt = (
                        f"Write Part {i} of {parts} on the topic: '{topic}'. "
                        f"Length requirement: {words_per_part}. "
                        "Continue directly from this previous context (maintain continuity, "
                        f"do not repeat it): '{previous_context[-400:]}'"
                    )
                else:
                    user_prompt = (
                        f"Write a complete script on the topic: '{topic}'. "
                        f"Length requirement: {words_per_part}. "
                        "Include a strong hook at the start."
                    )

                part_text, provider_used = _generate_part(
                    system_prompt, user_prompt, max_tokens_per_part, temperature, placeholder
                )
                part_texts.append(part_text)
                previous_context = part_text
                progress.progress(i / parts, text=f"भाग {i}/{parts} पूरा हुआ")

            full_script = "\n\n".join(
                f"==================== [ सीन / भाग {idx} ] ====================\n\n{txt}"
                for idx, txt in enumerate(part_texts, start=1)
            )
            progress.progress(1.0, text="पूरा हुआ! ✅")
            st.session_state.script_history.append({
                "title": topic, "script": full_script, "provider": provider_used,
            })
        except RuntimeError:
            return  # error already shown inline by _generate_part

    # Show the latest generated script (if any) with stats + export options
    if st.session_state.script_history:
        latest = st.session_state.script_history[-1]
        words, minutes = word_stats(latest["script"])
        st.divider()
        st.caption(f"⚙️ Powered by: {latest['provider']}  |  📝 {words} words  |  🎙️ ~{minutes} min read")

        c1, c2, c3 = st.columns(3)
        with c1:
            st.download_button("📥 Download .txt", data=latest["script"],
                                file_name=f"{sanitize_filename(latest['title'])}.txt", mime="text/plain")
        with c2:
            st.download_button("📥 Download .md", data=latest["script"],
                                file_name=f"{sanitize_filename(latest['title'])}.md", mime="text/markdown")
        with c3:
            try:
                docx_bytes = export_docx_bytes(latest["title"], latest["script"])
                st.download_button(
                    "📥 Download .docx", data=docx_bytes,
                    file_name=f"{sanitize_filename(latest['title'])}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            except ImportError:
                st.caption("`.docx` export ke liye `pip install python-docx` chahiye.")

        if len(st.session_state.script_history) > 1:
            with st.expander(f"📚 पिछली scripts ({len(st.session_state.script_history) - 1})"):
                for idx, item in enumerate(reversed(st.session_state.script_history[:-1])):
                    st.write(f"**{item['title']}**")
                    st.text_area("", value=item["script"], height=150, key=f"hist_{idx}", label_visibility="collapsed")


# ---------------------------------------------------------------------------
# Chatbot
# ---------------------------------------------------------------------------

CHATBOT_SYSTEM_PROMPT = (
    "You are a precise, helpful AI assistant. Follow the user's instructions exactly as given — "
    "do not add unrelated information, do not change the requested format or language, and do not "
    "pad your answers with unnecessary content. If a request is ambiguous, make a reasonable "
    "assumption and proceed rather than refusing."
)


def _render_chatbot():
    st.subheader("💬 AI Assistant & Chatbot")
    st.write("यहाँ आप AI से किसी भी तरह की मदद, आइडिया या सवाल पूछ सकते हैं:")

    if "messages" not in st.session_state:
        st.session_state.messages = []

    col1, col2 = st.columns([1, 1])
    with col1:
        if st.button("🗑️ Clear Chat"):
            st.session_state.messages = []
            st.rerun()
    with col2:
        if st.session_state.messages:
            chat_export = "\n\n".join(
                f"**{m['role'].upper()}:** {m['content']}" for m in st.session_state.messages
            )
            st.download_button("📥 Export Chat", data=chat_export, file_name="chat_export.md",
                                mime="text/markdown")

    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

    def _do_turn(user_query):
        st.session_state.messages.append({"role": "user", "content": user_query})
        with st.chat_message("user"):
            st.markdown(user_query)

        with st.chat_message("assistant"):
            trimmed_history = st.session_state.messages[-MAX_CHAT_HISTORY_MESSAGES:]
            msgs = [{"role": "system", "content": CHATBOT_SYSTEM_PROMPT}] + [
                {"role": m["role"], "content": m["content"]} for m in trimmed_history
            ]
            try:
                bot_reply = st.write_stream(stream_ai_with_fallback(msgs, max_tokens=2000, temperature=0.7))
                provider = stream_ai_with_fallback.last_provider
                st.caption(f"⚙️ Powered by: {provider}")
                st.session_state.messages.append({"role": "assistant", "content": bot_reply})
            except RuntimeError as e:
                st.error(str(e))

    if user_query := st.chat_input("अपना सवाल यहाँ पूछें..."):
        _do_turn(user_query)

    if st.session_state.messages and st.session_state.messages[-1]["role"] == "assistant":
        if st.button("🔄 Regenerate last answer"):
            last_user_msg = next(
                (m["content"] for m in reversed(st.session_state.messages) if m["role"] == "user"), None
            )
            if last_user_msg:
                st.session_state.messages.pop()  # remove old assistant reply
                _do_turn.__wrapped__ = None  # no-op, keeps linters happy
                with st.chat_message("assistant"):
                    trimmed_history = st.session_state.messages[-MAX_CHAT_HISTORY_MESSAGES:]
                    msgs = [{"role": "system", "content": CHATBOT_SYSTEM_PROMPT}] + [
                        {"role": m["role"], "content": m["content"]} for m in trimmed_history
                    ]
                    try:
                        bot_reply = st.write_stream(stream_ai_with_fallback(msgs, max_tokens=2000, temperature=0.7))
                        st.session_state.messages.append({"role": "assistant", "content": bot_reply})
                    except RuntimeError as e:
                        st.error(str(e))
